"""Service for exporting analysis results in various formats."""

import json
import csv
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List
import asyncio
from concurrent.futures import ThreadPoolExecutor

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.platypus import Image
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.barcharts import VerticalBarChart

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64

from app.models.schemas import AnalysisResult
from app.config import settings

class ExportService:
    """Service for exporting analysis results."""
    
    def __init__(self):
        """Initialize export service."""
        self.executor = ThreadPoolExecutor(max_workers=2)
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom styles for PDF generation."""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1a73e8'),
            spaceAfter=30,
            alignment=TA_CENTER
        ))
        
        # Heading style
        self.styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#1a73e8'),
            spaceAfter=12
        ))
        
        # Body style
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        ))
    
    async def export_result(self, result: AnalysisResult, format: str,
                          include_visualizations: bool = True,
                          include_detailed_feedback: bool = True) -> Path:
        """Export analysis result in specified format.
        
        Args:
            result: Analysis result to export
            format: Export format (pdf, csv, json, docx)
            include_visualizations: Whether to include charts
            include_detailed_feedback: Whether to include detailed feedback
            
        Returns:
            Path to exported file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"analysis_{timestamp}.{format}"
        file_path = settings.export_dir / filename
        
        if format == "pdf":
            await self._export_pdf(result, file_path, include_visualizations, include_detailed_feedback)
        elif format == "csv":
            await self._export_csv(result, file_path)
        elif format == "json":
            await self._export_json(result, file_path)
        elif format == "docx":
            await self._export_docx(result, file_path, include_visualizations, include_detailed_feedback)
        else:
            raise ValueError(f"Unsupported export format: {format}")
        
        return file_path
    
    async def _export_pdf(self, result: AnalysisResult, file_path: Path,
                         include_visualizations: bool, include_detailed_feedback: bool):
        """Export result as PDF."""
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            self.executor,
            self._generate_pdf,
            result, file_path, include_visualizations, include_detailed_feedback
        )
    
    def _generate_pdf(self, result: AnalysisResult, file_path: Path,
                     include_visualizations: bool, include_detailed_feedback: bool):
        """Generate PDF report."""
        doc = SimpleDocTemplate(
            str(file_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        story = []
        
        # Title
        story.append(Paragraph("Text Analysis Report", self.styles['CustomTitle']))
        story.append(Spacer(1, 0.2*inch))
        
        # Summary section
        story.append(Paragraph("Executive Summary", self.styles['CustomHeading']))
        story.append(Paragraph(result.feedback_summary, self.styles['CustomBody']))
        story.append(Spacer(1, 0.2*inch))
        
        # Overall score
        score_data = [
            ['Overall Score', f"{result.overall_score:.1f}/100"],
            ['Grammar Score', f"{result.grammar.score:.1f}/100"],
            ['Coherence Score', f"{result.coherence.score:.1f}/100"],
            ['Relevance Score', f"{result.relevance.score:.1f}/100"]
        ]
        
        score_table = Table(score_data, colWidths=[3*inch, 2*inch])
        score_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(score_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Text statistics
        story.append(Paragraph("Text Statistics", self.styles['CustomHeading']))
        stats_data = [
            ['Word Count', str(result.word_count)],
            ['Sentence Count', str(result.sentence_count)],
            ['Paragraph Count', str(result.paragraph_count)],
            ['Avg. Sentence Length', f"{result.avg_sentence_length:.1f} words"]
        ]
        
        stats_table = Table(stats_data, colWidths=[3*inch, 2*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(stats_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Visualizations
        if include_visualizations:
            story.append(Paragraph("Score Visualization", self.styles['CustomHeading']))
            
            # Create pie chart
            drawing = Drawing(400, 200)
            pie = Pie()
            pie.x = 150
            pie.y = 50
            pie.width = 100
            pie.height = 100
            pie.data = [result.grammar.score, result.coherence.score, result.relevance.score]
            pie.labels = ['Grammar', 'Coherence', 'Relevance']
            pie.slices.strokeWidth = 0.5
            pie.slices[0].fillColor = colors.HexColor('#4285f4')
            pie.slices[1].fillColor = colors.HexColor('#34a853')
            pie.slices[2].fillColor = colors.HexColor('#fbbc04')
            drawing.add(pie)
            
            story.append(drawing)
            story.append(Spacer(1, 0.3*inch))
        
        # Strengths and improvements
        story.append(Paragraph("Strengths", self.styles['CustomHeading']))
        for strength in result.strengths:
            story.append(Paragraph(f"• {strength}", self.styles['CustomBody']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("Areas for Improvement", self.styles['CustomHeading']))
        for improvement in result.areas_for_improvement:
            story.append(Paragraph(f"• {improvement}", self.styles['CustomBody']))
        story.append(Spacer(1, 0.2*inch))
        
        # Detailed feedback
        if include_detailed_feedback:
            story.append(PageBreak())
            story.append(Paragraph("Detailed Analysis", self.styles['CustomTitle']))
            story.append(Spacer(1, 0.3*inch))
            
            # Grammar details
            story.append(Paragraph("Grammar Analysis", self.styles['CustomHeading']))
            story.append(Paragraph(f"Error Count: {len(result.grammar.errors)}", self.styles['CustomBody']))
            story.append(Paragraph(f"Error Density: {result.grammar.error_density:.2f}%", self.styles['CustomBody']))
            
            if result.grammar.errors[:5]:  # Show first 5 errors
                story.append(Paragraph("Sample Errors:", self.styles['CustomBody']))
                for error in result.grammar.errors[:5]:
                    story.append(Paragraph(
                        f"• {error.message} (Position: {error.position[0]}-{error.position[1]})",
                        self.styles['CustomBody']
                    ))
            
            story.append(Spacer(1, 0.2*inch))
            
            # Coherence details
            story.append(Paragraph("Coherence Analysis", self.styles['CustomHeading']))
            if result.coherence.weak_connections:
                story.append(Paragraph(
                    f"Weak Connections Found: {len(result.coherence.weak_connections)}",
                    self.styles['CustomBody']
                ))
            
            # Readability scores
            if result.coherence.readability_scores:
                story.append(Paragraph("Readability Metrics:", self.styles['CustomBody']))
                for metric, score in list(result.coherence.readability_scores.items())[:4]:
                    story.append(Paragraph(
                        f"• {metric.replace('_', ' ').title()}: {score:.1f}",
                        self.styles['CustomBody']
                    ))
        
        # Build PDF
        doc.build(story)
    
    async def _export_csv(self, result: AnalysisResult, file_path: Path):
        """Export result as CSV."""
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            self.executor,
            self._generate_csv,
            result, file_path
        )
    
    def _generate_csv(self, result: AnalysisResult, file_path: Path):
        """Generate CSV file."""
        with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            
            # Header
            writer.writerow(['Metric', 'Value'])
            
            # Scores
            writer.writerow(['Overall Score', result.overall_score])
            writer.writerow(['Grammar Score', result.grammar.score])
            writer.writerow(['Coherence Score', result.coherence.score])
            writer.writerow(['Relevance Score', result.relevance.score])
            
            # Statistics
            writer.writerow(['Word Count', result.word_count])
            writer.writerow(['Sentence Count', result.sentence_count])
            writer.writerow(['Paragraph Count', result.paragraph_count])
            writer.writerow(['Average Sentence Length', result.avg_sentence_length])
            
            # Processing time
            writer.writerow(['Processing Time (seconds)', result.processing_time])
            
            # Timestamp
            writer.writerow(['Analysis Date', result.timestamp.isoformat()])
            
            # Feedback
            writer.writerow([''])
            writer.writerow(['Feedback Summary'])
            writer.writerow([result.feedback_summary])
            
            # Strengths
            writer.writerow([''])
            writer.writerow(['Strengths'])
            for strength in result.strengths:
                writer.writerow([strength])
            
            # Improvements
            writer.writerow([''])
            writer.writerow(['Areas for Improvement'])
            for improvement in result.areas_for_improvement:
                writer.writerow([improvement])
            
            # Grammar errors summary
            writer.writerow([''])
            writer.writerow(['Grammar Error Summary'])
            writer.writerow(['Error Type', 'Count'])
            if hasattr(result.grammar, 'details') and 'error_counts' in result.grammar.details:
                for error_type, count in result.grammar.details['error_counts'].items():
                    writer.writerow([error_type, count])
    
    async def _export_json(self, result: AnalysisResult, file_path: Path):
        """Export result as JSON."""
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            self.executor,
            self._generate_json,
            result, file_path
        )
    
    def _generate_json(self, result: AnalysisResult, file_path: Path):
        """Generate JSON file."""
        # Convert result to dict
        result_dict = result.dict()
        
        # Add metadata
        result_dict['export_metadata'] = {
            'export_date': datetime.now().isoformat(),
            'version': '1.0.0',
            'format': 'json'
        }
        
        # Write JSON
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(result_dict, f, indent=2, ensure_ascii=False, default=str)
    
    async def _export_docx(self, result: AnalysisResult, file_path: Path,
                          include_visualizations: bool, include_detailed_feedback: bool):
        """Export result as DOCX."""
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(
            self.executor,
            self._generate_docx,
            result, file_path, include_visualizations, include_detailed_feedback
        )
    
    def _generate_docx(self, result: AnalysisResult, file_path: Path,
                      include_visualizations: bool, include_detailed_feedback: bool):
        """Generate DOCX file."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Text Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(result.feedback_summary)
        doc.add_paragraph()
        
        # Scores section
        doc.add_heading('Analysis Scores', level=1)
        
        # Create scores table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Score'
        
        # Data rows
        score_data = [
            ('Overall Score', f'{result.overall_score:.1f}/100'),
            ('Grammar Score', f'{result.grammar.score:.1f}/100'),
            ('Coherence Score', f'{result.coherence.score:.1f}/100'),
            ('Relevance Score', f'{result.relevance.score:.1f}/100')
        ]
        
        for i, (metric, score) in enumerate(score_data, 1):
            cells = table.rows[i].cells
            cells[0].text = metric
            cells[1].text = score
        
        doc.add_paragraph()
        
        # Text Statistics
        doc.add_heading('Text Statistics', level=1)
        stats_paragraph = doc.add_paragraph()
        stats_paragraph.add_run(f'Word Count: ').bold = True
        stats_paragraph.add_run(f'{result.word_count}\n')
        stats_paragraph.add_run(f'Sentence Count: ').bold = True
        stats_paragraph.add_run(f'{result.sentence_count}\n')
        stats_paragraph.add_run(f'Paragraph Count: ').bold = True
        stats_paragraph.add_run(f'{result.paragraph_count}\n')
        stats_paragraph.add_run(f'Average Sentence Length: ').bold = True
        stats_paragraph.add_run(f'{result.avg_sentence_length:.1f} words')
        
        doc.add_paragraph()
        
        # Strengths
        doc.add_heading('Strengths', level=1)
        for strength in result.strengths:
            doc.add_paragraph(f'• {strength}', style='List Bullet')
        
        doc.add_paragraph()
        
        # Areas for Improvement
        doc.add_heading('Areas for Improvement', level=1)
        for improvement in result.areas_for_improvement:
            doc.add_paragraph(f'• {improvement}', style='List Bullet')
        
        # Detailed feedback
        if include_detailed_feedback:
            doc.add_page_break()
            doc.add_heading('Detailed Analysis', level=1)
            
            # Grammar details
            doc.add_heading('Grammar Analysis', level=2)
            doc.add_paragraph(f'Total Errors Found: {len(result.grammar.errors)}')
            doc.add_paragraph(f'Error Density: {result.grammar.error_density:.2f}%')
            
            if result.grammar.suggestions:
                doc.add_heading('Grammar Suggestions:', level=3)
                for suggestion in result.grammar.suggestions:
                    doc.add_paragraph(f'• {suggestion}', style='List Bullet')
            
            # Coherence details
            doc.add_heading('Coherence Analysis', level=2)
            if result.coherence.weak_connections:
                doc.add_paragraph(f'Weak Connections Found: {len(result.coherence.weak_connections)}')
            
            if result.coherence.suggestions:
                doc.add_heading('Coherence Suggestions:', level=3)
                for suggestion in result.coherence.suggestions:
                    doc.add_paragraph(f'• {suggestion}', style='List Bullet')
            
            # Relevance details
            if result.relevance.topic_coverage:
                doc.add_heading('Topic Coverage', level=2)
                for topic, coverage in result.relevance.topic_coverage.items():
                    doc.add_paragraph(f'{topic}: {coverage:.1f}%')
        
        # Save document
        doc.save(str(file_path))