"""Text preprocessing utilities."""

import re
from typing import Optional, List
import unicodedata
from pathlib import Path
import PyPDF2
from docx import Document
import chardet
import logging

logger = logging.getLogger(__name__)

class TextPreprocessor:
    """Utility class for text preprocessing."""
    
    def __init__(self):
        """Initialize preprocessor."""
        self.logger = logger
    
    def process(self, text: str) -> str:
        """Process text for analysis.
        
        Args:
            text: Raw text input
            
        Returns:
            Processed text
        """
        # Remove null bytes
        text = text.replace('\x00', '')
        
        # Normalize unicode
        text = unicodedata.normalize('NFKD', text)
        
        # Fix common encoding issues
        text = self._fix_encoding_issues(text)
        
        # Normalize whitespace
        text = self._normalize_whitespace(text)
        
        # Fix quotes and apostrophes
        text = self._normalize_quotes(text)
        
        # Remove control characters
        text = self._remove_control_characters(text)
        
        # Ensure proper sentence spacing
        text = self._fix_sentence_spacing(text)
        
        return text.strip()
    
    async def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """Extract text from various file formats.
        
        Args:
            file_content: File content as bytes
            filename: Original filename with extension
            
        Returns:
            Extracted text
        """
        extension = Path(filename).suffix.lower()
        
        if extension == '.txt':
            return self._extract_from_txt(file_content)
        elif extension == '.pdf':
            return self._extract_from_pdf(file_content)
        elif extension in ['.docx', '.doc']:
            return self._extract_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _extract_from_txt(self, content: bytes) -> str:
        """Extract text from TXT file.
        
        Args:
            content: File content
            
        Returns:
            Extracted text
        """
        # Detect encoding
        detected = chardet.detect(content)
        encoding = detected['encoding'] or 'utf-8'
        
        try:
            text = content.decode(encoding)
        except UnicodeDecodeError:
            # Fallback to utf-8 with error handling
            text = content.decode('utf-8', errors='ignore')
            self.logger.warning("Had to ignore some characters during decoding")
        
        return self.process(text)
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF file.
        
        Args:
            content: PDF file content
            
        Returns:
            Extracted text
        """
        import io
        
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text_parts.append(page.extract_text())
        
        text = '\n'.join(text_parts)
        return self.process(text)
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX file.
        
        Args:
            content: DOCX file content
            
        Returns:
            Extracted text
        """
        import io
        
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        text = '\n\n'.join(text_parts)
        return self.process(text)
    
    def _fix_encoding_issues(self, text: str) -> str:
        """Fix common encoding issues.
        
        Args:
            text: Input text
            
        Returns:
            Fixed text
        """
        # Common replacements
        replacements = {
            'â€™': "'",
            'â€œ': '"',
            'â€': '"',
            'â€"': '—',
            'â€"': '–',
            'Ã©': 'é',
            'Ã¨': 'è',
            'Ã ': 'à',
            'Ã¢': 'â',
            'Ã´': 'ô',
            'Ã®': 'î',
            'Ã§': 'ç',
            'Ã‰': 'É',
            'âˆ'': '-',
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        return text
    
    def _normalize_whitespace(self, text: str) -> str:
        """Normalize whitespace in text.
        
        Args:
            text: Input text
            
        Returns:
            Text with normalized whitespace
        """
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Replace multiple newlines with double newline (paragraph break)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove spaces before punctuation
        text = re.sub(r' +([.,!?;:])', r'\1', text)
        
        # Ensure space after punctuation
        text = re.sub(r'([.,!?;:])([A-Za-z])', r'\1 \2', text)
        
        return text
    
    def _normalize_quotes(self, text: str) -> str:
        """Normalize quotes and apostrophes.
        
        Args:
            text: Input text
            
        Returns:
            Text with normalized quotes
        """
        # Smart quotes to regular quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        # Fix apostrophes
        text = re.sub(r'(\w)'(\w)', r"\1'\2", text)  # Contractions
        text = re.sub(r'(\w)'s\b', r"\1's", text)    # Possessives
        
        return text
    
    def _remove_control_characters(self, text: str) -> str:
        """Remove control characters from text.
        
        Args:
            text: Input text
            
        Returns:
            Cleaned text
        """
        # Remove all control characters except newlines and tabs
        control_chars = ''.join(
            chr(i) for i in range(32) 
            if chr(i) not in ['\n', '\t', '\r']
        )
        translator = str.maketrans('', '', control_chars)
        
        return text.translate(translator)
    
    def _fix_sentence_spacing(self, text: str) -> str:
        """Ensure proper spacing between sentences.
        
        Args:
            text: Input text
            
        Returns:
            Text with fixed sentence spacing
        """
        # Ensure single space after sentence-ending punctuation
        text = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', text)
        
        # Fix missing spaces after commas
        text = re.sub(r',([A-Za-z])', r', \1', text)
        
        return text
    
    def split_into_chunks(self, text: str, chunk_size: int = 1000, 
                         overlap: int = 100) -> List[str]:
        """Split text into overlapping chunks for processing.
        
        Args:
            text: Text to split
            chunk_size: Size of each chunk in characters
            overlap: Number of overlapping characters
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Find end position
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence end
                sentence_end = text.rfind('.', start, end)
                if sentence_end > start + chunk_size // 2:
                    end = sentence_end + 1
            
            chunks.append(text[start:end])
            
            # Move start position
            start = end - overlap
            
            # Ensure we don't create tiny final chunks
            if len(text) - start < chunk_size // 2 and chunks:
                # Merge with previous chunk
                chunks[-1] = chunks[-1] + text[start:]
                break
        
        return chunks